from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# profile - picture
document.add_picture('suresh cahn.jpg', width=Inches(2.0))

# name, number, email
name = input('what is your name ? ')
speak('Hello' + name + 'how are you today?')
speak('ni number kottu')
phone_number = input('what is your phone number?')
speak('your number is ' + phone_number)
email = input('what is your email?')
document.add_paragraph(name + ' | ' + phone_number + ' | ' + email)
# about me
document.add_heading('About Me')
about_me = input("Tell me about yourself ?")
document.add_paragraph(about_me)

# WORK EXPERIENCE
document.add_heading('Work experience')
p = document.add_paragraph()
p.style = 'List Bullet'
company = input('enter company')
from_date = input('From Date')
to_date = input('To Date')
p.add_run(company + ' ').bold = True
p.add_run(from_date+'-' + to_date + '\n')

details = input('Describe your details at' + company)
p.add_run(details)

document.save('cv.docx')
